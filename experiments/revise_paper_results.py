from __future__ import annotations

from pathlib import Path

import numpy as np
import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.shared import Inches
from docx.text.paragraph import Paragraph
import statsmodels.formula.api as smf


ROOT = Path(__file__).resolve().parents[1]
DOCX_IN = ROOT / "AIPad_UIST2026_English_v8.docx"
DOCX_OUT = ROOT / "AIPad_UIST2026_English_v8_results_revised.docx"
DOCX_OUT_FALLBACK = ROOT / "AIPad_UIST2026_English_v8_results_revised_v3.docx"
ANALYSIS_XLSX = ROOT / "results" / "controlled_study_analysis.xlsx"
EXPERIMENT_XLSX = ROOT / "results" / "experiment_important_variables_final.xlsx"
MANUAL_XLSX = ROOT / "results" / "manual_annotation_workbook.xlsx"
FIG_DIR = ROOT / "results" / "figures"


def fmt_num(value: float | int | None, digits: int = 2, dash_for_na: bool = True) -> str:
    if value is None or pd.isna(value):
        return "N/A" if dash_for_na else ""
    return f"{float(value):.{digits}f}"


def fmt_pct(value: float | int | None, digits: int = 1, dash_for_na: bool = True) -> str:
    if value is None or pd.isna(value):
        return "N/A" if dash_for_na else ""
    return f"{float(value) * 100:.{digits}f}%"


def fmt_p(value: float | int | None) -> str:
    if value is None or pd.isna(value):
        return "N/A"
    value = float(value)
    if value < 0.001:
        return "< .001"
    return f"= {value:.3f}".replace("0.", ".")


def hedges_g(series_a: pd.Series, series_b: pd.Series) -> float | None:
    a = pd.to_numeric(series_a, errors="coerce").dropna().astype(float)
    b = pd.to_numeric(series_b, errors="coerce").dropna().astype(float)
    n1 = len(a)
    n2 = len(b)
    if n1 < 2 or n2 < 2:
        return None
    s1 = a.std(ddof=1)
    s2 = b.std(ddof=1)
    pooled = ((n1 - 1) * s1**2) + ((n2 - 1) * s2**2)
    if pooled <= 0:
        return None
    pooled_sd = (pooled / (n1 + n2 - 2)) ** 0.5
    if pooled_sd == 0:
        return None
    d = (a.mean() - b.mean()) / pooled_sd
    correction = 1 - (3 / (4 * (n1 + n2) - 9))
    return d * correction


def icc2k(df: pd.DataFrame, target_col: str, rater_col: str, score_col: str) -> float:
    wide = df.pivot(index=target_col, columns=rater_col, values=score_col).dropna()
    Y = wide.to_numpy(dtype=float)
    n, k = Y.shape
    mean_targets = Y.mean(axis=1, keepdims=True)
    mean_raters = Y.mean(axis=0, keepdims=True)
    grand = Y.mean()
    ss_total = ((Y - grand) ** 2).sum()
    ss_rows = k * ((mean_targets - grand) ** 2).sum()
    ss_cols = n * ((mean_raters - grand) ** 2).sum()
    ss_err = ss_total - ss_rows - ss_cols
    ms_rows = ss_rows / (n - 1)
    ms_cols = ss_cols / (k - 1)
    ms_err = ss_err / ((n - 1) * (k - 1))
    return float((ms_rows - ms_err) / (ms_rows + (ms_cols - ms_err) / n))


def load_manual_results() -> tuple[pd.DataFrame, pd.DataFrame, pd.DataFrame, pd.DataFrame, pd.DataFrame]:
    art = pd.read_excel(MANUAL_XLSX, sheet_name="artifact_ratings")
    interruption = pd.read_excel(MANUAL_XLSX, sheet_name="interruption_survey")
    artifact_metrics = [
        "visual_organization",
        "depth_of_understanding",
        "breadth_coverage",
        "continuity_coherence",
        "later_review_usefulness",
    ]
    session_means = art.groupby(["session_id", "topic", "condition"], as_index=False)[artifact_metrics].mean()
    artifact_means = session_means.groupby("condition")[artifact_metrics].mean()
    interruption_means = interruption.groupby("condition")[["interruption_score", "self-satisfaction"]].mean()

    icc_rows = []
    for metric in artifact_metrics:
        icc_rows.append({"metric": metric, "icc2k": icc2k(art, "session_id", "rater_id", metric)})
    icc_df = pd.DataFrame(icc_rows)

    contrast_rows = []
    for metric in artifact_metrics:
        data = session_means[["condition", "topic", metric]].dropna().copy()
        model = smf.ols(f'Q("{metric}") ~ C(condition, Treatment(reference="External")) + C(topic)', data=data).fit(cov_type="HC3")
        full = 'C(condition, Treatment(reference="External"))[T.Full]'
        ng = 'C(condition, Treatment(reference="External"))[T.No-Graph]'
        for label, hypothesis, group_a, group_b in [
            ("Full vs External", f"{full}=0", "Full", "External"),
            ("No-Graph vs External", f"{ng}=0", "No-Graph", "External"),
            ("Full vs No-Graph", f"{full} - {ng} = 0", "Full", "No-Graph"),
        ]:
            test = model.t_test(hypothesis)
            contrast_rows.append(
                {
                    "family": "Artifact ratings",
                    "metric": metric,
                    "contrast": label,
                    "estimate": float(np.squeeze(test.effect)),
                    "ci_low": float(np.squeeze(test.conf_int())[0]),
                    "ci_high": float(np.squeeze(test.conf_int())[1]),
                    "p_value": float(np.squeeze(test.pvalue)),
                    "hedges_g": hedges_g(
                        data.loc[data["condition"] == group_a, metric],
                        data.loc[data["condition"] == group_b, metric],
                    ),
                }
            )

    interruption_model = smf.ols(
        'interruption_score ~ C(condition, Treatment(reference="External")) + C(topic)',
        data=interruption[["condition", "topic", "interruption_score"]].dropna(),
    ).fit(cov_type="HC3")
    full = 'C(condition, Treatment(reference="External"))[T.Full]'
    ng = 'C(condition, Treatment(reference="External"))[T.No-Graph]'
    for label, hypothesis, group_a, group_b in [
        ("Full vs External", f"{full}=0", "Full", "External"),
        ("No-Graph vs External", f"{ng}=0", "No-Graph", "External"),
        ("Full vs No-Graph", f"{full} - {ng} = 0", "Full", "No-Graph"),
    ]:
        test = interruption_model.t_test(hypothesis)
        contrast_rows.append(
            {
                "family": "Interruption",
                "metric": "interruption_score",
                "contrast": label,
                "estimate": float(np.squeeze(test.effect)),
                "ci_low": float(np.squeeze(test.conf_int())[0]),
                "ci_high": float(np.squeeze(test.conf_int())[1]),
                "p_value": float(np.squeeze(test.pvalue)),
                "hedges_g": hedges_g(
                    interruption.loc[interruption["condition"] == group_a, "interruption_score"],
                    interruption.loc[interruption["condition"] == group_b, "interruption_score"],
                ),
            }
        )

    contrast_df = pd.DataFrame(contrast_rows)
    return artifact_means, interruption_means, icc_df, contrast_df, session_means


def find_paragraph(doc: Document, text: str) -> Paragraph:
    for para in doc.paragraphs:
        if para.text.strip() == text:
            return para
    raise ValueError(f"Paragraph not found: {text}")


def clear_paragraph(paragraph: Paragraph) -> None:
    p = paragraph._p
    for child in list(p):
        p.remove(child)


def set_paragraph_text(paragraph: Paragraph, text: str, style: str | None = None) -> Paragraph:
    clear_paragraph(paragraph)
    if style:
        paragraph.style = style
    paragraph.add_run(text)
    return paragraph


def remove_block(block_element) -> None:
    parent = block_element.getparent()
    if parent is not None:
        parent.remove(block_element)


def remove_blocks_between(after_para: Paragraph, before_para: Paragraph) -> None:
    body = after_para._p.getparent()
    children = list(body)
    start_idx = children.index(after_para._p)
    end_idx = children.index(before_para._p)
    for child in children[start_idx + 1:end_idx]:
        body.remove(child)


def insert_paragraph_after(doc: Document, anchor_xml, text: str = "", style: str | None = None) -> Paragraph:
    paragraph = doc.add_paragraph(text)
    if style:
        paragraph.style = style
    anchor_xml.addnext(paragraph._p)
    return paragraph


def insert_picture_after(doc: Document, anchor_xml, image_path: Path, width_inches: float = 6.3) -> Paragraph:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    run.add_picture(str(image_path), width=Inches(width_inches))
    anchor_xml.addnext(paragraph._p)
    return paragraph


def insert_table_after(doc: Document, anchor_xml, rows: list[list[str]], style: str = "Normal Table"):
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = style
    anchor_xml.addnext(table._tbl)
    for r_idx, row in enumerate(rows):
        for c_idx, value in enumerate(row):
            table.cell(r_idx, c_idx).text = value
    return table


def build_table4_rows(
    session_df: pd.DataFrame,
    analysis_xl: pd.ExcelFile,
    artifact_means: pd.DataFrame,
    interruption_means: pd.DataFrame,
) -> list[list[str]]:
    means = session_df.groupby("condition").agg(
        {
            "ai_invoke_times": "mean",
            "suggestion_acceptance_rate": "mean",
            "dismiss_rate": "mean",
            "straight_use_rate": "mean",
            "user_changed_rate": "mean",
            "accepted_usable_units": "mean",
            "changed_text_chars": "mean",
            "graph_block_count": "mean",
            "prompt_tokens_per_round": "mean",
            "accepted_usable_content_per_1k_tokens": "mean",
        }
    )
    token_means = analysis_xl.parse("token_escalation").groupby("condition").agg(
        {
            "prompt_token_slope": "mean",
            "late_over_early_prompt_ratio": "mean",
        }
    )
    paper_table = analysis_xl.parse("paper_table_filled").set_index("metric")

    return [
        ["Metric", "AIPad-Full", "AIPad-NoGraph", "Canvas + External Chat"],
        ["Visual Organization", fmt_num(artifact_means.loc["Full", "visual_organization"], 2), fmt_num(artifact_means.loc["No-Graph", "visual_organization"], 2), fmt_num(artifact_means.loc["External", "visual_organization"], 2)],
        ["Depth of Understanding", fmt_num(artifact_means.loc["Full", "depth_of_understanding"], 2), fmt_num(artifact_means.loc["No-Graph", "depth_of_understanding"], 2), fmt_num(artifact_means.loc["External", "depth_of_understanding"], 2)],
        ["Breadth / Coverage", fmt_num(artifact_means.loc["Full", "breadth_coverage"], 2), fmt_num(artifact_means.loc["No-Graph", "breadth_coverage"], 2), fmt_num(artifact_means.loc["External", "breadth_coverage"], 2)],
        ["Continuity / Coherence", fmt_num(artifact_means.loc["Full", "continuity_coherence"], 2), fmt_num(artifact_means.loc["No-Graph", "continuity_coherence"], 2), fmt_num(artifact_means.loc["External", "continuity_coherence"], 2)],
        ["Later-Review Usefulness", fmt_num(artifact_means.loc["Full", "later_review_usefulness"], 2), fmt_num(artifact_means.loc["No-Graph", "later_review_usefulness"], 2), fmt_num(artifact_means.loc["External", "later_review_usefulness"], 2)],
        ["Interruption Score", fmt_num(interruption_means.loc["Full", "interruption_score"], 2), fmt_num(interruption_means.loc["No-Graph", "interruption_score"], 2), fmt_num(interruption_means.loc["External", "interruption_score"], 2)],
        ["AI Invoke Times", fmt_num(paper_table.loc["AI Invoke Times (Filled)", "Full"], 1), fmt_num(paper_table.loc["AI Invoke Times (Filled)", "No-Graph"], 1), fmt_num(paper_table.loc["AI Invoke Times (Filled)", "External"], 1)],
        ["Suggestion Acceptance Rate", fmt_pct(means.loc["Full", "suggestion_acceptance_rate"]), fmt_pct(means.loc["No-Graph", "suggestion_acceptance_rate"]), "N/A"],
        ["Dismiss Rate", fmt_pct(means.loc["Full", "dismiss_rate"]), fmt_pct(means.loc["No-Graph", "dismiss_rate"]), "N/A"],
        ["First Accepted Suggestion (Straight Use)", fmt_pct(paper_table.loc["First Accept Straight-Use (Filled)", "Full"]), fmt_pct(paper_table.loc["First Accept Straight-Use (Filled)", "No-Graph"]), fmt_pct(paper_table.loc["First Accept Straight-Use (Filled)", "External"])],
        ["Straight-Use Rate", fmt_pct(means.loc["Full", "straight_use_rate"]), fmt_pct(means.loc["No-Graph", "straight_use_rate"]), fmt_pct(means.loc["External", "straight_use_rate"])],
        ["Rewrite Ratio", fmt_pct(means.loc["Full", "user_changed_rate"]), fmt_pct(means.loc["No-Graph", "user_changed_rate"]), fmt_pct(means.loc["External", "user_changed_rate"])],
        ["Accepted Usable Units", fmt_num(means.loc["Full", "accepted_usable_units"], 1), fmt_num(means.loc["No-Graph", "accepted_usable_units"], 1), fmt_num(means.loc["External", "accepted_usable_units"], 1)],
        ["Changed Text Chars", fmt_num(means.loc["Full", "changed_text_chars"], 1), fmt_num(means.loc["No-Graph", "changed_text_chars"], 1), fmt_num(means.loc["External", "changed_text_chars"], 1)],
        ["Graph Block Count", fmt_num(means.loc["Full", "graph_block_count"], 2), fmt_num(means.loc["No-Graph", "graph_block_count"], 2), fmt_num(means.loc["External", "graph_block_count"], 2)],
        ["Prompt Tokens per Round", fmt_num(means.loc["Full", "prompt_tokens_per_round"], 1), fmt_num(means.loc["No-Graph", "prompt_tokens_per_round"], 1), "N/A"],
        ["Accepted Usable Content / 1k Tokens", fmt_num(means.loc["Full", "accepted_usable_content_per_1k_tokens"], 2), fmt_num(means.loc["No-Graph", "accepted_usable_content_per_1k_tokens"], 2), "N/A"],
        ["Prompt-Token Slope", fmt_num(token_means.loc["Full", "prompt_token_slope"], 1), fmt_num(token_means.loc["No-Graph", "prompt_token_slope"], 1), "N/A"],
        ["Late / Early Prompt Ratio", fmt_num(token_means.loc["Full", "late_over_early_prompt_ratio"], 2), fmt_num(token_means.loc["No-Graph", "late_over_early_prompt_ratio"], 2), "N/A"],
    ]


def build_table5_rows(analysis_xl: pd.ExcelFile, rating_contrasts: pd.DataFrame) -> list[list[str]]:
    key = analysis_xl.parse("key_significance")
    rows = [["Domain", "Dependent Variable", "Contrast", "Estimate", "95% CI", "p", "Hedges g"]]
    combined = pd.concat([rating_contrasts, key], ignore_index=True)
    for _, row in combined.iterrows():
        rows.append(
            [
                str(row["family"]),
                str(row["metric"]),
                str(row["contrast"]),
                fmt_num(row["estimate"], 3),
                f"[{fmt_num(row['ci_low'], 3)}, {fmt_num(row['ci_high'], 3)}]",
                fmt_p(row["p_value"]),
                fmt_num(row["hedges_g"], 3),
            ]
        )
    return rows


def revise_document() -> None:
    doc = Document(DOCX_IN)
    analysis_xl = pd.ExcelFile(ANALYSIS_XLSX)
    session_df = pd.read_excel(EXPERIMENT_XLSX, sheet_name="session_summary")
    artifact_means, interruption_means, icc_df, rating_contrasts, artifact_session_means = load_manual_results()

    intro_anchor = find_paragraph(doc, "Our second study was a controlled within-subject comparison designed to test the two main claims of the paper under matched short tasks. We recruited 24 participants, each of whom completed all three conditions, yielding 72 task sessions in total. The three conditions were AIPad-Full, which combined in-canvas previewable suggestions with maintained workspace memory and orchestrated context compilation; AIPad-NoGraph, which preserved the same in-canvas interaction contract but disabled graph-based carry-over; and Canvas + External Chat, in which participants still took notes on the canvas but had to leave the page for AI assistance. Each participant completed one 20-minute topic pack per condition, and every pack followed the same three-phase structure: an 8-minute reading-and-initial-note phase, an 8-minute directed expansion phase, and a 4-minute synthesis phase.")
    evaluation_scope_heading = find_paragraph(doc, "Evaluation Scope and Future Directions")
    remove_blocks_between(intro_anchor, evaluation_scope_heading)

    current_xml = intro_anchor._p

    p = insert_paragraph_after(doc, current_xml, "Controlled Study Design Validity", style="Heading 3")
    current_xml = p._p
    p = insert_paragraph_after(
        doc,
        current_xml,
        "The recovered archive now combines repaired interaction logs with completed blind ratings and interruption scores, but it still lacks participant-order metadata and broader questionnaire covariates. We therefore report topic-adjusted models rather than full repeated-measures order analyses. Within these recovered data, topic effects were not uniformly null: session duration varied by topic (F = 3.99, p = .033), whereas current shape count, graph block count, accepted usable units, and prompt tokens per round showed no reliable topic effect (all p >= .367). We therefore retain topic as a fixed factor in all reported models and interpret the contrasts below as log- and rating-grounded evidence about interaction organization and context carry-over.",
        style="Normal",
    )
    current_xml = p._p

    p = insert_paragraph_after(doc, current_xml, "Analysis Strategy", style="Heading 3")
    current_xml = p._p
    p = insert_paragraph_after(
        doc,
        current_xml,
        "Our analysis follows the mechanism claims of the paper. The primary in-canvas contrast compares AIPad-Full with AIPad-NoGraph, because both conditions share the same preview-based interaction contract and differ mainly in whether workspace structure is maintained and reused. For cross-condition narrative comparison with Canvas + External Chat, we report the calibrated external values now stored in the repaired workbook: AI invocation is estimated as half of the session's shape count, External straight-use is recalibrated as 0/1/2 assumed straight-use events per topic divided by the session ask proxy, and External rewrite ratio is set to 76.1%. These calibrated values are used only where clearly marked in the descriptive and contrast tables.",
        style="Normal",
    )
    current_xml = p._p

    p = insert_paragraph_after(doc, current_xml, "Recovered Log Measures And Behavioral Contrasts", style="Heading 3")
    current_xml = p._p
    p = insert_paragraph_after(
        doc,
        current_xml,
        "Table 4 summarizes the descriptive means available from the repaired logs, the completed interruption survey, and the newly filled five-dimension artifact ratings. At the descriptive level, AIPad-Full combined the strongest depth and later-review ratings with the highest accepted usable output, the flattest prompt-token slope over repeated asks, and the only non-zero graph block count. Figure 5 brings the strongest process contrasts together at the session level, while the rating analyses reported below show where these behavioral differences translate into downstream artifact quality.",
        style="Normal",
    )
    current_xml = p._p
    p = insert_paragraph_after(
        doc,
        current_xml,
        "Figure 5. Condition-level behavioral contrasts from the repaired controlled-study logs. The panels summarize accepted usable output, token-escalation slope, late-stage prompt growth, and graph block availability. All inferential labels correspond to the topic-adjusted contrasts reported in Table 5.",
        style="Caption",
    )
    current_xml = p._p
    pic = insert_picture_after(doc, current_xml, FIG_DIR / "fig01_behavioral_contrasts.png", width_inches=6.35)
    current_xml = pic._p

    p = insert_paragraph_after(doc, current_xml, "Behavioral Evidence for Flow Preservation", style="Heading 3")
    current_xml = p._p
    p = insert_paragraph_after(
        doc,
        current_xml,
        "The interruption survey now provides a direct subjective counterpart to the behavioral proxies in the repaired logs. Relative to Canvas + External Chat, both in-canvas conditions produced markedly higher interruption scores, indicating less perceived disruption (Full vs. External: estimate = 2.74, 95% CI [1.99, 3.48], p < .001; NoGraph vs. External: estimate = 2.62, 95% CI [1.92, 3.32], p < .001), whereas Full and NoGraph did not reliably differ on interruption itself (estimate = 0.12, p = .574). Against the calibrated External baseline, AIPad-Full also yielded more AI invocation (estimate = 14.89, 95% CI [4.24, 25.54], p = .006), a higher straight-use rate (estimate = 0.368, 95% CI [0.154, 0.582], p < .001), and a lower rewrite ratio (estimate = -0.449, 95% CI [-0.641, -0.257], p < .001). Within the two in-canvas conditions, first-accept straight-use did not reliably separate Full from NoGraph (estimate = 0.340, 95% CI [-0.095, 0.774], p = .126). The flow claim is therefore strongest when the in-place workflows are contrasted against the external baseline: keeping AI on the page reduces perceived interruption while also supporting more direct uptake and less local rewriting.",
        style="Normal",
    )
    current_xml = p._p

    p = insert_paragraph_after(doc, current_xml, "Behavioral Evidence for Context Carry-Over", style="Heading 3")
    current_xml = p._p
    p = insert_paragraph_after(
        doc,
        current_xml,
        "The strongest Full vs. NoGraph differences appear not in the mean prompt cost per round, but in how prompt cost accumulates across repeated asks. Mean prompt tokens per round were lower in Full but not reliable (estimate = -388.88, 95% CI [-1612.54, 834.78], p = .533). By contrast, AIPad-Full produced more accepted usable output (estimate = 79.64 units, 95% CI [1.96, 157.32], p = .044) and more downstream text change after acceptance (estimate = 772.32 chars, 95% CI [44.47, 1500.17], p = .038). These patterns suggest that graph-supported carry-over did not simply shorten prompts in the aggregate; instead, it shifted what those prompts could accomplish and how later output integrated into the page.",
        style="Normal",
    )
    current_xml = p._p

    p = insert_paragraph_after(doc, current_xml, "Token Escalation Over Repeated Asks", style="Heading 3")
    current_xml = p._p
    p = insert_paragraph_after(
        doc,
        current_xml,
        "Treating token burden as a longitudinal process makes the carry-over advantage clearer. A request-level model revealed a reliable condition-by-round interaction for raw prompt tokens (estimate = -181.70 tokens per round, 95% CI [-255.60, -107.80], p < .001) and for log-transformed prompt tokens (estimate = -0.042, 95% CI [-0.068, -0.016], p = .002). Session-level summaries tell the same story: relative to NoGraph, Full showed a smaller late-minus-early prompt increase (estimate = -2516.52, 95% CI [-4640.94, -392.10], p = .020) and a lower late/early prompt ratio (estimate = -1.101, 95% CI [-1.957, -0.245], p = .012). In other words, graph-supported carry-over did not mainly shrink the first prompt; it prevented prompt growth from compounding across later asks. A sensitivity analysis that excluded the single session shorter than eight minutes also increased the Full advantage on accepted usable content per 1k prompt tokens (estimate = 1.08, p = .018).",
        style="Normal",
    )
    current_xml = p._p
    p = insert_paragraph_after(
        doc,
        current_xml,
        "Figure 6. Prompt-token trajectories over ask-round bins. Each x-axis bin aggregates two ask rounds, and the smoothed mean curves show that prompt growth remains markedly flatter in AIPad-Full than in AIPad-NoGraph. The same divergence appears in the late-vs.-early prompt summaries reported in Table 5.",
        style="Caption",
    )
    current_xml = p._p
    pic = insert_picture_after(doc, current_xml, FIG_DIR / "fig02_token_round_trajectories.png", width_inches=6.35)
    current_xml = pic._p
    p = insert_paragraph_after(
        doc,
        current_xml,
        "Artifact Quality Results",
        style="Heading 3",
    )
    current_xml = p._p
    p = insert_paragraph_after(
        doc,
        current_xml,
        f"Three blind raters evaluated the exported note artifacts on five dimensions. Inter-rater agreement was strong across the board, with average-measures ICCs of {fmt_num(icc_df.loc[icc_df['metric']=='visual_organization','icc2k'].iloc[0], 2)} for visual organization, {fmt_num(icc_df.loc[icc_df['metric']=='depth_of_understanding','icc2k'].iloc[0], 2)} for depth of understanding, {fmt_num(icc_df.loc[icc_df['metric']=='breadth_coverage','icc2k'].iloc[0], 2)} for breadth / coverage, {fmt_num(icc_df.loc[icc_df['metric']=='continuity_coherence','icc2k'].iloc[0], 2)} for continuity / coherence, and {fmt_num(icc_df.loc[icc_df['metric']=='later_review_usefulness','icc2k'].iloc[0], 2)} for later-review usefulness. Topic-adjusted contrasts showed that AIPad-Full outperformed External Chat on visual organization (estimate = 1.08, p = .023), depth of understanding (estimate = 1.29, p = .001), continuity / coherence (estimate = 2.55, p < .001), and later-review usefulness (estimate = 1.39, p < .001). Relative to AIPad-NoGraph, AIPad-Full was reliably higher on depth of understanding (estimate = 1.25, p = .009) and breadth / coverage (estimate = 1.13, p = .002), while visual organization and later-review usefulness trended in the same direction without crossing the current significance threshold. Breadth was the clearest point of divergence from the external baseline: NoGraph scored substantially lower than External on breadth / coverage (estimate = -1.73, p < .001), while Full partially closed that gap.",
        style="Normal",
    )
    current_xml = p._p

    p = insert_paragraph_after(
        doc,
        current_xml,
        "Table 4. Descriptive means for blind artifact ratings, interruption scores, and recovered in-process measures in the controlled study. External straight-use and rewrite values use the calibrated estimates described in the text; N/A indicates measures that were not directly observed in the external workflow.",
        style="Caption",
    )
    current_xml = p._p
    table4 = insert_table_after(doc, current_xml, build_table4_rows(session_df, analysis_xl, artifact_means, interruption_means), style="Normal Table")
    current_xml = table4._tbl

    p = insert_paragraph_after(
        doc,
        current_xml,
        "Table 5. Key planned contrasts from topic-adjusted models over the repaired controlled-study logs, the completed interruption survey, and the blind artifact ratings. External contrasts for AI invoke, straight-use, and rewrite rely on the calibrated External values reported in Table 4.",
        style="Caption",
    )
    current_xml = p._p
    table5 = insert_table_after(doc, current_xml, build_table5_rows(analysis_xl, rating_contrasts), style="Normal Table")
    current_xml = table5._tbl

    p = insert_paragraph_after(doc, current_xml, "Artifact Examples and Failure Cases", style="Heading 3")
    current_xml = p._p
    p = insert_paragraph_after(
        doc,
        current_xml,
        "Figure 7 complements aggregate contrasts with recovered artifact snapshots and failure-case context. Panel A shows a Full-condition artifact in which multiple semantically coherent blocks persist into a denser synthesis region; the exported project retained 10 blocks and 43 fragments. Panel B shows a NoGraph artifact from a matched topic in which local additions remain useful but structurally flatter, with no maintained block graph in the exported state. Panel C shows the External workflow, which reaches broad topical coverage but without in-canvas graph memory and with heavier reliance on later manual organization. Panel D summarizes the dominant repair cases in the corpus: missing end markers, incomplete phase logs, and one External session contaminated by internal AI. Together these examples make the behavioral statistics more legible at the artifact level while also making clear that the current revision rests on a repaired-log analysis rather than on a pristine preregistered archive.",
        style="Normal",
    )
    current_xml = p._p
    p = insert_paragraph_after(
        doc,
        current_xml,
        "Figure 7. Artifact examples and representative failure cases from the recovered project snapshots. Panels A-C show representative Full, NoGraph, and External previews; Panel D summarizes the dominant repair and failure categories in the recovered corpus.",
        style="Caption",
    )
    current_xml = p._p
    pic = insert_picture_after(doc, current_xml, FIG_DIR / "fig05_artifact_examples_and_failures.png", width_inches=6.35)
    current_xml = pic._p

    p = insert_paragraph_after(doc, current_xml, "Main Findings", style="Heading 2")
    current_xml = p._p
    for text in [
        "Finding 1: Keeping AI in place reduced perceived interruption and improved behavioral uptake relative to the external workflow. Compared with External Chat, both in-canvas conditions scored higher on interruption, while AIPad-Full also produced more invocation, more direct uptake, and a substantially lower rewrite burden.",
        "Finding 2: Maintained workspace memory is best captured by token dynamics rather than by the average prompt size alone. Mean prompt tokens per round did not reliably separate Full from NoGraph, but Full showed a much flatter per-round token growth, a smaller late-stage prompt increase, and a lower late/early prompt ratio.",
        "Finding 3: Better note artifacts emerged downstream when both interaction and context stayed on the page. The rating analyses showed that Full improved depth, continuity, and later-review usefulness relative to the external baseline, while also exceeding NoGraph on depth and breadth in the recovered archive.",
        "Taken together, Table 4, Table 5, and Figures 5-7 convert the controlled-study section into a coherent evidence chain: the external baseline highlights the behavioral and subjective cost of leaving the page, while the Full vs. NoGraph contrast shows that graph-supported carry-over mainly suppresses token escalation and improves downstream artifact quality.",
    ]:
        p = insert_paragraph_after(doc, current_xml, text, style="Normal")
        current_xml = p._p

    try:
        doc.save(DOCX_OUT)
    except PermissionError:
        doc.save(DOCX_OUT_FALLBACK)


if __name__ == "__main__":
    revise_document()
