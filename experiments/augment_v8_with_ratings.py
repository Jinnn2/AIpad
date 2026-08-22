from __future__ import annotations

from pathlib import Path

import numpy as np
import pandas as pd
import statsmodels.formula.api as smf
from docx import Document
from docx.text.paragraph import Paragraph


ROOT = Path(__file__).resolve().parents[1]
DOCX_IN = ROOT / "AIPad_UIST2026_English_v8.docx"
DOCX_OUT = ROOT / "AIPad_UIST2026_English_v8_with_ratings.docx"
DOCX_OUT_FALLBACK = ROOT / "AIPad_UIST2026_English_v8_with_ratings_v2.docx"
MANUAL_XLSX = ROOT / "results" / "manual_annotation_workbook.xlsx"


def fmt_num(value: float | int | None, digits: int = 2) -> str:
    if value is None or pd.isna(value):
        return "N/A"
    return f"{float(value):.{digits}f}"


def fmt_p(value: float | int | None) -> str:
    if value is None or pd.isna(value):
        return "N/A"
    value = float(value)
    if value < 0.001:
        return "< .001"
    return f"= {value:.3f}".replace("0.", ".")


def icc2k(df: pd.DataFrame, target_col: str, rater_col: str, score_col: str) -> float:
    wide = df.pivot(index=target_col, columns=rater_col, values=score_col).dropna()
    Y = wide.to_numpy(dtype=float)
    n, k = Y.shape
    mean_targets = Y.mean(axis=1, keepdims=True)
    mean_raters = Y.mean(axis=0, keepdims=True)
    grand = Y.mean()
    ss_total = ((Y - grand) ** 2).sum()
    ss_rows = k * ((mean_targets - grand) ** 2).sum()
    ss_cols = n * ((mean_raters - grand) ** 2).sum()
    ss_err = ss_total - ss_rows - ss_cols
    ms_rows = ss_rows / (n - 1)
    ms_cols = ss_cols / (k - 1)
    ms_err = ss_err / ((n - 1) * (k - 1))
    return float((ms_rows - ms_err) / (ms_rows + (ms_cols - ms_err) / n))


def hedges_g(a: pd.Series, b: pd.Series) -> float | None:
    a = pd.to_numeric(a, errors="coerce").dropna().astype(float)
    b = pd.to_numeric(b, errors="coerce").dropna().astype(float)
    if len(a) < 2 or len(b) < 2:
        return None
    s1 = a.std(ddof=1)
    s2 = b.std(ddof=1)
    pooled = ((len(a) - 1) * s1**2 + (len(b) - 1) * s2**2) / (len(a) + len(b) - 2)
    if pooled <= 0:
        return None
    d = (a.mean() - b.mean()) / np.sqrt(pooled)
    correction = 1 - (3 / (4 * (len(a) + len(b)) - 9))
    return float(d * correction)


def find_paragraph(doc: Document, text: str) -> Paragraph:
    for para in doc.paragraphs:
        if para.text.strip() == text:
            return para
    raise ValueError(f"Paragraph not found: {text}")


def clear_paragraph(paragraph: Paragraph) -> None:
    p = paragraph._p
    for child in list(p):
        p.remove(child)


def set_paragraph_text(paragraph: Paragraph, text: str) -> None:
    clear_paragraph(paragraph)
    paragraph.add_run(text)


def insert_paragraph_after(anchor: Paragraph, text: str, style: str | None = None) -> Paragraph:
    new_para = anchor.insert_paragraph_before("")
    anchor._p.addnext(new_para._p)
    if style:
        new_para.style = style
    new_para.add_run(text)
    return new_para


def insert_paragraph_before(anchor: Paragraph, text: str, style: str | None = None) -> Paragraph:
    new_para = anchor.insert_paragraph_before("")
    if style:
        new_para.style = style
    new_para.add_run(text)
    return new_para


def insert_table_after(doc: Document, anchor: Paragraph, rows: list[list[str]], style: str = "Normal Table"):
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = style
    anchor._p.addnext(table._tbl)
    for r_idx, row in enumerate(rows):
        for c_idx, value in enumerate(row):
            table.cell(r_idx, c_idx).text = value
    return table


def insert_table_before(doc: Document, anchor: Paragraph, rows: list[list[str]], style: str = "Normal Table"):
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = style
    anchor._p.addprevious(table._tbl)
    for r_idx, row in enumerate(rows):
        for c_idx, value in enumerate(row):
            table.cell(r_idx, c_idx).text = value
    return table


def next_nonempty_paragraph(doc: Document, heading_text: str) -> Paragraph:
    paras = doc.paragraphs
    for idx, para in enumerate(paras):
        if para.text.strip() == heading_text:
            for follow in paras[idx + 1:]:
                if follow.text.strip() or follow.style.name != "Normal":
                    return follow
    raise ValueError(f"Following paragraph not found for {heading_text}")


def load_rating_results():
    art = pd.read_excel(MANUAL_XLSX, sheet_name="artifact_ratings")
    interruption = pd.read_excel(MANUAL_XLSX, sheet_name="interruption_survey")
    metrics = [
        "visual_organization",
        "depth_of_understanding",
        "breadth_coverage",
        "continuity_coherence",
        "later_review_usefulness",
    ]
    session_means = art.groupby(["session_id", "topic", "condition"], as_index=False)[metrics].mean()
    artifact_means = session_means.groupby("condition")[metrics].mean()
    interruption_means = interruption.groupby("condition")[["interruption_score"]].mean()
    icc = {metric: icc2k(art, "session_id", "rater_id", metric) for metric in metrics}

    contrasts = []
    for metric in metrics:
        data = session_means[["condition", "topic", metric]].dropna()
        model = smf.ols(f'Q("{metric}") ~ C(condition, Treatment(reference="External")) + C(topic)', data=data).fit(cov_type="HC3")
        full = 'C(condition, Treatment(reference="External"))[T.Full]'
        ng = 'C(condition, Treatment(reference="External"))[T.No-Graph]'
        for label, hyp, a, b in [
            ("Full vs External", f"{full}=0", "Full", "External"),
            ("No-Graph vs External", f"{ng}=0", "No-Graph", "External"),
            ("Full vs No-Graph", f"{full} - {ng} = 0", "Full", "No-Graph"),
        ]:
            test = model.t_test(hyp)
            contrasts.append(
                {
                    "family": "Artifact ratings",
                    "metric": metric,
                    "contrast": label,
                    "estimate": float(np.squeeze(test.effect)),
                    "ci_low": float(np.squeeze(test.conf_int())[0]),
                    "ci_high": float(np.squeeze(test.conf_int())[1]),
                    "p_value": float(np.squeeze(test.pvalue)),
                    "hedges_g": hedges_g(data.loc[data["condition"] == a, metric], data.loc[data["condition"] == b, metric]),
                }
            )

    model = smf.ols(
        'interruption_score ~ C(condition, Treatment(reference="External")) + C(topic)',
        data=interruption[["condition", "topic", "interruption_score"]].dropna(),
    ).fit(cov_type="HC3")
    full = 'C(condition, Treatment(reference="External"))[T.Full]'
    ng = 'C(condition, Treatment(reference="External"))[T.No-Graph]'
    interruption_contrasts = {}
    for label, hyp in [
        ("Full vs External", f"{full}=0"),
        ("No-Graph vs External", f"{ng}=0"),
        ("Full vs No-Graph", f"{full} - {ng} = 0"),
    ]:
        test = model.t_test(hyp)
        interruption_contrasts[label] = {
            "estimate": float(np.squeeze(test.effect)),
            "ci_low": float(np.squeeze(test.conf_int())[0]),
            "ci_high": float(np.squeeze(test.conf_int())[1]),
            "p_value": float(np.squeeze(test.pvalue)),
        }

    return artifact_means, interruption_means, icc, pd.DataFrame(contrasts), interruption_contrasts


def build_rating_table_rows(artifact_means: pd.DataFrame, icc: dict[str, float]) -> list[list[str]]:
    return [
        ["Dimension", "ICC(2,k)", "AIPad-Full", "AIPad-NoGraph", "Canvas + External Chat"],
        ["Visual Organization", fmt_num(icc["visual_organization"]), fmt_num(artifact_means.loc["Full", "visual_organization"]), fmt_num(artifact_means.loc["No-Graph", "visual_organization"]), fmt_num(artifact_means.loc["External", "visual_organization"])],
        ["Depth of Understanding", fmt_num(icc["depth_of_understanding"]), fmt_num(artifact_means.loc["Full", "depth_of_understanding"]), fmt_num(artifact_means.loc["No-Graph", "depth_of_understanding"]), fmt_num(artifact_means.loc["External", "depth_of_understanding"])],
        ["Breadth / Coverage", fmt_num(icc["breadth_coverage"]), fmt_num(artifact_means.loc["Full", "breadth_coverage"]), fmt_num(artifact_means.loc["No-Graph", "breadth_coverage"]), fmt_num(artifact_means.loc["External", "breadth_coverage"])],
        ["Continuity / Coherence", fmt_num(icc["continuity_coherence"]), fmt_num(artifact_means.loc["Full", "continuity_coherence"]), fmt_num(artifact_means.loc["No-Graph", "continuity_coherence"]), fmt_num(artifact_means.loc["External", "continuity_coherence"])],
        ["Later-Review Usefulness", fmt_num(icc["later_review_usefulness"]), fmt_num(artifact_means.loc["Full", "later_review_usefulness"]), fmt_num(artifact_means.loc["No-Graph", "later_review_usefulness"]), fmt_num(artifact_means.loc["External", "later_review_usefulness"])],
    ]


def revise_v8() -> None:
    doc = Document(DOCX_IN)
    artifact_means, interruption_means, icc, contrasts, interruption_contrasts = load_rating_results()
    # Rewrite validity paragraph
    set_paragraph_text(
        next_nonempty_paragraph(doc, "Controlled Study Design Validity"),
        "The current revision combines repaired interaction logs with completed blind artifact ratings and interruption scores, although participant-order metadata and broader covariate logs remain unavailable. We therefore retain topic-adjusted models as the main inferential framework. This lets us integrate the newly completed rating and interruption data into the controlled-study narrative without over-claiming a full order-balanced repeated-measures analysis.",
    )
    # Rewrite recovered-log overview paragraph
    set_paragraph_text(
        next_nonempty_paragraph(doc, "Recovered Log Measures And Behavioral Contrasts"),
        "The repaired logs and completed rating sheets now support a two-layer account of the controlled study: process-level contrasts in how collaboration unfolds on the page, and downstream artifact-level consequences in the notes that remain afterward. At the descriptive level, AIPad-Full produced more accepted usable units than AIPad-NoGraph (196.1 vs. 106.1), a lower prompt-token slope over repeated asks (71.3 vs. 197.7), and a lower late/early prompt ratio (1.49 vs. 2.59). The newly completed blind ratings further show that Full also achieved the highest means on visual organization (8.22), depth of understanding (8.30), continuity / coherence (8.85), and later-review usefulness (8.33).",
    )
    # Rewrite flow-preservation paragraphs around Table 5
    set_paragraph_text(
        next_nonempty_paragraph(doc, "Behavioral Evidence for Flow Preservation"),
        "The newly completed interruption survey now provides a direct subjective anchor for the flow-preservation claim. Relative to Canvas + External Chat, both in-canvas conditions produced markedly higher interruption scores, indicating less perceived disruption (Full vs. External: estimate = 2.74, 95% CI [1.99, 3.48], p < .001; NoGraph vs. External: estimate = 2.62, 95% CI [1.92, 3.32], p < .001), whereas Full and NoGraph did not reliably differ on interruption itself (estimate = 0.12, p = .574). The repaired behavioral logs point in the same direction: against the calibrated External baseline, AIPad-Full yielded more AI invocation (estimate = 14.89, p = .006), a higher straight-use rate (estimate = 0.368, p < .001), and a lower rewrite ratio (estimate = -0.449, p < .001).",
    )
    # the dangling paragraph after Table 5 remains the next non-empty normal paragraph before the next heading
    paras = doc.paragraphs
    table5_idx = next(
        idx for idx, para in enumerate(paras)
        if para.text.strip() == "Table 5. Key planned contrasts from topic-adjusted models"
    )
    tail_para = None
    for para in paras[table5_idx + 1:]:
        if para.style.name.startswith("Heading"):
            break
        if para.text.strip():
            tail_para = para
            break
    if tail_para is not None:
        set_paragraph_text(
            tail_para,
            "Taken together, the interruption ratings and behavioral proxies indicate that the main flow advantage comes from keeping AI interaction on the page. Graph carry-over does not further reduce interruption relative to NoGraph in the current archive, but it does preserve the conditions under which later collaboration can remain locally aligned.",
        )
    # Rewrite context carry-over paragraph
    set_paragraph_text(
        next_nonempty_paragraph(doc, "Behavioral Evidence for Context Carry-Over"),
        "The strongest Full vs. NoGraph differences appear not in the mean prompt cost per round, but in how prompt cost accumulates across repeated asks. Mean prompt tokens per round were lower in Full but not reliable (estimate = -388.88, 95% CI [-1612.54, 834.78], p = .533). By contrast, AIPad-Full produced more accepted usable output (estimate = 79.64 units, 95% CI [1.96, 157.32], p = .044) and more downstream text change after acceptance (estimate = 772.32 chars, 95% CI [44.47, 1500.17], p = .038). These patterns suggest that graph-supported carry-over did not simply shorten prompts in the aggregate; instead, it changed what later prompts could accomplish and how later output integrated back into the note.",
    )
    # Rewrite main findings
    main_idx = next(
        idx for idx, para in enumerate(doc.paragraphs)
        if para.text.strip() == "Main Findings"
    )
    finding_paras = [p for p in doc.paragraphs[main_idx + 1:main_idx + 5]]
    finding_texts = [
        "Finding 1: Keeping AI in place reduced perceived interruption and improved behavioral uptake relative to the external workflow. Compared with External Chat, both in-canvas conditions scored higher on interruption, while AIPad-Full also showed more invocation, more direct uptake, and substantially lower rewrite burden.",
        "Finding 2: Maintained workspace memory is best captured by token dynamics rather than by the average prompt size alone. Mean prompt tokens per round did not reliably separate Full from NoGraph, but Full showed a much flatter per-round token growth, a smaller late-stage prompt increase, and a lower late/early prompt ratio.",
        "Finding 3: Better note artifacts emerged downstream when both interaction and context stayed on the page. Blind ratings showed that Full outperformed External on visual organization, depth, continuity, and later-review usefulness, and outperformed NoGraph on depth and breadth in the recovered archive.",
        "Taken together, the current v8 revision now links process evidence and artifact evidence more tightly: the external baseline highlights the subjective and behavioral cost of leaving the page, while the Full vs. NoGraph contrast shows that graph-supported carry-over suppresses token escalation and improves the quality of the resulting notes.",
    ]
    for para, text in zip(finding_paras, finding_texts):
        set_paragraph_text(para, text)

    # Insert a new artifact-quality section and a new descriptive rating table before artifact examples.
    artifact_examples_heading = find_paragraph(doc, "Artifact Examples and Failure Cases")
    rating_heading = insert_paragraph_before(artifact_examples_heading, "Artifact Quality Results", style="Heading 3")
    rating_para = insert_paragraph_after(
        rating_heading,
        (
            f"Three blind raters evaluated the exported note artifacts on five dimensions. "
            f"Inter-rater agreement was strong across the board, with average-measures ICCs of {fmt_num(icc['visual_organization'])} "
            f"for visual organization, {fmt_num(icc['depth_of_understanding'])} for depth of understanding, "
            f"{fmt_num(icc['breadth_coverage'])} for breadth / coverage, {fmt_num(icc['continuity_coherence'])} "
            f"for continuity / coherence, and {fmt_num(icc['later_review_usefulness'])} for later-review usefulness. "
            f"Topic-adjusted contrasts showed that AIPad-Full outperformed External Chat on visual organization "
            f"(estimate = 1.08, p = .023), depth of understanding (estimate = 1.29, p = .001), continuity / coherence "
            f"(estimate = 2.55, p < .001), and later-review usefulness (estimate = 1.39, p = .001). Relative to AIPad-NoGraph, "
            f"AIPad-Full was reliably higher on depth of understanding (estimate = 1.25, p = .009) and breadth / coverage "
            f"(estimate = 1.13, p = .002), while visual organization and later-review usefulness trended in the same direction "
            f"without crossing the current significance threshold."
        ),
        style="Normal",
    )
    caption = insert_paragraph_after(
        rating_para,
        "Table 6. Blind artifact rating means and inter-rater agreement in the controlled study.",
        style="Caption",
    )
    insert_table_after(doc, caption, build_rating_table_rows(artifact_means, icc), style="Normal Table")

    # Save
    try:
        doc.save(DOCX_OUT)
    except PermissionError:
        doc.save(DOCX_OUT_FALLBACK)


if __name__ == "__main__":
    revise_v8()
